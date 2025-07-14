from docx import Document
from docx.shared import Inches
from io import BytesIO

# --- Charter Generator ---
def generate_charter_docx(project_name, data):
    """Generates a professionally formatted Project Charter DOCX."""
    doc = Document()
    doc.add_heading(f"Project Charter: {project_name}", 0)

    # Helper to add sections
    def add_section(title, content):
        doc.add_heading(title, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(str(content))
        doc.add_paragraph() # Add a space after the section

    add_section("Project Description", data.get("project_description", "N/A"))
    add_section("Key Objectives", data.get("objectives", []))
    add_section("High-Level Requirements", data.get("requirements", []))
    add_section("Key Stakeholders", data.get("stakeholders", []))
    add_section("Budget", f"${data.get('budget', 0):,}")
    add_section("Timeline", f"{data.get('timeline_weeks', 0)} weeks")
    
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- Scope Generator ---
def generate_scope_docx(project_name, data):
    """Generates a professionally formatted Scope Statement DOCX."""
    doc = Document()
    doc.add_heading(f"Scope Statement: {project_name}", 0)

    def add_section(title, content):
        doc.add_heading(title, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(str(content))
        doc.add_paragraph()

    add_section("Project Scope Description", data.get("description", "N/A"))
    add_section("Key Deliverables", data.get("deliverables", []))
    add_section("Acceptance Criteria", data.get("acceptance_criteria", []))
    add_section("Exclusions", data.get("exclusions", []))
    add_section("Constraints", data.get("constraints", []))
    add_section("Assumptions", data.get("assumptions", []))

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- WBS Generator ---
def generate_wbs_docx(project_name, data):
    """Generates a professionally formatted WBS DOCX with a nested structure."""
    doc = Document()
    doc.add_heading(f"Work Breakdown Structure: {project_name}", 0)
    doc.add_paragraph(
        "This document breaks down the project deliverables into smaller, more manageable components."
    )
    
    # Recursive helper function to add tasks and sub-tasks with indentation
    def add_tasks_recursively(tasks, level=0):
        for task in tasks:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * level)
            p.add_run(task.get("task_name", "Unnamed Task"))
            
            # If there are sub-tasks, recurse
            if "sub_tasks" in task and task["sub_tasks"]:
                add_tasks_recursively(task["sub_tasks"], level + 1)

    wbs_items = data.get("wbs_items", [])
    add_tasks_recursively(wbs_items)

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- Risk Register Generator ---
def generate_risk_register_docx(project_name, data):
    """Generates a professionally formatted Risk Register DOCX with a table."""
    doc = Document()
    doc.add_heading(f"Risk Register: {project_name}", 0)
    doc.add_paragraph(
        "This document identifies potential risks, their assessment, and planned response strategies."
    )
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "Risk Description"
    header_cells[1].text = "Probability"
    header_cells[2].text = "Impact"
    header_cells[3].text = "Response Strategy"
    header_cells[4].text = "Owner"
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # THE FIX: Ensure 'risks' is a list, even if the key is missing or its value is None.
    risks = data.get("risks") or []
    
    for risk in risks:
        # Ensure each risk is a dictionary to prevent further errors
        if not isinstance(risk, dict):
            continue # Skip malformed risk entries

        row_cells = table.add_row().cells
        row_cells[0].text = risk.get("risk_description", "")
        row_cells[1].text = risk.get("probability", "")
        row_cells[2].text = risk.get("impact", "")
        row_cells[3].text = risk.get("response_strategy", "")
        row_cells[4].text = risk.get("owner", "")
        
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def generate_stakeholder_analysis_docx(data: dict, **kwargs):
    """Generates a .docx file for the Stakeholder Analysis."""
    document = Document()
    # You could optionally use the project_name from kwargs here if needed
    # project_name = kwargs.get("project_name", "Project")
    document.add_heading('Stakeholder Analysis', level=1)
    
    stakeholders = data.get('stakeholders', [])
    if not stakeholders:
        document.add_paragraph("No stakeholder data available.")
    else:
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Role'
        hdr_cells[2].text = 'Interest'
        hdr_cells[3].text = 'Influence'
        hdr_cells[4].text = 'Engagement Strategy'
        
        for item in stakeholders:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('name', '')
            row_cells[1].text = item.get('role', '')
            row_cells[2].text = item.get('interest', '')
            row_cells[3].text = item.get('influence', '')
            row_cells[4].text = item.get('engagement_strategy', '')
            
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_communication_plan_docx(data: dict, **kwargs):
    """Generates a .docx file for the Communication Plan."""
    document = Document()
    document.add_heading('Project Communication Plan', level=1)
    
    communications = data.get('communications', [])
    if not communications:
        document.add_paragraph("No communication plan data available.")
    else:
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Stakeholder'
        hdr_cells[1].text = 'Information'
        hdr_cells[2].text = 'Method'
        hdr_cells[3].text = 'Frequency'
        hdr_cells[4].text = 'Owner'
        
        for item in communications:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('stakeholder', '')
            row_cells[1].text = item.get('information', '')
            row_cells[2].text = item.get('method', '')
            row_cells[3].text = item.get('frequency', '')
            row_cells[4].text = item.get('owner', '')

    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io